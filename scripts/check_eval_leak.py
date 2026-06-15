import json
from collections.abc import Iterator
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from genacademy_coach.corpus import iter_indexable_files, load_corpus_document
from genacademy_coach.eval_split import normalized_words, phrase_hashes
from genacademy_coach.settings import CoachSettings

SCAN_GLOBS = [
    "AGENTS.md",
    "README.md",
    "docs/**/*.md",
    "specs/**/*.md",
    "src/**/*.py",
    "scripts/**/*.py",
]


def read_eval_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs]
        table_cells = [
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        ]
        return "\n".join([*paragraphs, *table_cells])
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    return ""


def normalized_text(text: str) -> str:
    return " ".join(normalized_words(text))


def iter_committed_scan_texts(settings: CoachSettings) -> Iterator[tuple[Path, str]]:
    for pattern in SCAN_GLOBS:
        for path in settings.repo_root.glob(pattern):
            if path.is_file():
                yield path, path.read_text(encoding="utf-8", errors="ignore")


def iter_local_corpus_scan_texts(settings: CoachSettings) -> Iterator[tuple[Path, str]]:
    for path in iter_indexable_files(settings.corpus_dir):
        yield path, load_corpus_document(path).text


def main() -> None:
    settings = CoachSettings.from_env()
    if not settings.eval_manifest_path.exists():
        raise SystemExit(f"missing eval manifest: {settings.eval_manifest_path}")
    manifest = json.loads(settings.eval_manifest_path.read_text(encoding="utf-8"))
    test_items = [item for item in manifest["items"] if item["split"] == "test"]
    needles = {item["id"] for item in test_items} | {item["source_sha256"] for item in test_items}
    offenders: list[str] = []
    eval_phrase_sources: list[tuple[str, str]] = []
    missing_eval_sources: list[str] = []
    for item in test_items:
        eval_path = settings.eval_questions_dir / item["source_file"]
        if not eval_path.exists():
            missing_eval_sources.append(item["source_file"])
            continue
        eval_phrase_sources.append((item["source_file"], read_eval_text(eval_path)))
    test_phrase_hashes = phrase_hashes(eval_phrase_sources)
    scan_texts = list(iter_committed_scan_texts(settings))

    # This direct scan is deliberately simple and adequate for current corpus size. If the corpus
    # grows to thousands of files, replace it with trie/Aho-Corasick matching before making
    # it CI-hard.
    if test_phrase_hashes:
        scan_texts.extend(iter_local_corpus_scan_texts(settings))
    for path, text in scan_texts:
        if any(needle in text for needle in needles):
            offenders.append(str(path))
        normalized = normalized_text(text)
        for phrase, matches in test_phrase_hashes.items():
            if phrase in normalized:
                phrase_refs = ", ".join(
                    f"{match['source_file']}:{match['phrase_hash']}" for match in matches
                )
                offenders.append(f"{path} matched eval phrase {phrase_refs}")
    if offenders:
        raise SystemExit("eval test leak detected in: " + ", ".join(sorted(set(offenders))))
    if missing_eval_sources:
        print(
            "private eval sources missing; skipped local n-gram leak scan for: "
            + ", ".join(sorted(missing_eval_sources))
        )
    print(
        "no eval test IDs/checksums found in code/docs; no eval n-grams found where "
        "private eval sources were available "
        f"({len(test_items)} test items)"
    )


if __name__ == "__main__":
    main()
