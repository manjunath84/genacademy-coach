from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


def read_eval_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs]
        table_cells = [
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        ]
        return "\n".join([*paragraphs, *table_cells])
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    return ""
